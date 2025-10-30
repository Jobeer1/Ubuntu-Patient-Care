"""
Script to create a Philips Radiography QC Checklist Word document (Test 46-49)
with proper tables and checkboxes
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_checkbox(cell):
    """Add a checkbox to a table cell"""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    # Add checkbox symbol
    run.text = '☐'
    run.font.size = Pt(14)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def shade_cell(cell, color):
    """Shade a cell with a specific color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_philips_qc_checklist_2():
    """Create the Philips QC Checklist Word document (Tests 46-49)"""
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width = Inches(11)  # Landscape
        section.page_height = Inches(8.5)
    
    # Add Philips logo text (right-aligned)
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_run = logo_para.add_run('PHILIPS')
    logo_run.font.size = Pt(16)
    logo_run.font.bold = True
    logo_run.font.color.rgb = RGBColor(0, 0, 139)
    
    doc.add_paragraph()  # Space
    
    # Create table: 2 label columns + 31 day columns = 33 columns
    # Rows: header (2) + Test 46 items (6) + Test 47 items (2) + Test 48 items (2) + Test 49 (1) + Comments (1) = 14 rows
    table = doc.add_table(rows=14, cols=33)
    table.style = 'Table Grid'
    
    # Row 0: Month and Year, DoH Licence #, Unit Serial #
    cell = table.cell(0, 0)
    cell.text = 'Month and Year:'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    shade_cell(cell, 'F0F0F0')
    
    # Merge cells for Month and Year input
    merged_month = table.cell(0, 1).merge(table.cell(0, 10))
    
    # DoH Licence #
    cell = table.cell(0, 11)
    cell.text = 'DoH Licence #:'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    shade_cell(cell, 'F0F0F0')
    merged_doh = table.cell(0, 12).merge(table.cell(0, 21))
    
    # Unit Serial #
    cell = table.cell(0, 22)
    cell.text = 'Unit Serial #:'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    shade_cell(cell, 'F0F0F0')
    merged_serial = table.cell(0, 23).merge(table.cell(0, 32))
    
    # Row 1: Date headers (1-31)
    date_cell = table.cell(1, 0)
    date_cell.text = 'Date:'
    date_cell.paragraphs[0].runs[0].font.bold = True
    date_cell.paragraphs[0].runs[0].font.size = Pt(10)
    shade_cell(date_cell, 'F0F0F0')
    
    # Empty cell for alignment
    table.cell(1, 1).text = ''
    
    for day in range(1, 32):
        cell = table.cell(1, day + 1)
        cell.text = str(day)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(cell, 'F0F0F0')
    
    # Test 46 items (rows 2-7)
    test_46_items = [
        ('Radiographer Initials:', False),
        ('Lead Aprons:', False),
        ('Radiation Warning Lights:', False),
        ('Radiation Warning Sign:', False),
        ('Indicators:', False),
        ('Movement Controls:', False),
        ('Display Monitors Cleanliness:', False),
        ('Flicker:', False),
    ]
    
    # Row 2: Test 46 label
    cell = table.cell(2, 0)
    cell.text = 'Test 46:'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    shade_cell(cell, 'E0E0E0')
    
    # Radiographer Initials
    cell = table.cell(2, 1)
    cell.text = 'Radiographer Initials:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(2, day + 1))
    
    # Lead Aprons
    cell = table.cell(3, 1)
    cell.text = 'Lead Aprons:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(3, day + 1))
    
    # Radiation Warning Lights
    cell = table.cell(4, 1)
    cell.text = 'Radiation Warning Lights:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(4, day + 1))
    
    # Radiation Warning Sign
    cell = table.cell(5, 1)
    cell.text = 'Radiation Warning Sign:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(5, day + 1))
    
    # Indicators
    cell = table.cell(6, 1)
    cell.text = 'Indicators:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(6, day + 1))
    
    # Movement Controls
    cell = table.cell(7, 1)
    cell.text = 'Movement Controls:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(7, day + 1))
    
    # Display Monitors Cleanliness
    cell = table.cell(8, 1)
    cell.text = 'Display Monitors Cleanliness:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(8, day + 1))
    
    # Flicker
    cell = table.cell(9, 1)
    cell.text = 'Flicker:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(9, day + 1))
    
    # Merge Test 46 cells vertically (rows 2-9)
    merged_test46 = table.cell(2, 0).merge(table.cell(9, 0))
    merged_test46.text = 'Test 46:'
    merged_test46.paragraphs[0].runs[0].font.bold = True
    merged_test46.paragraphs[0].runs[0].font.size = Pt(10)
    merged_test46.vertical_alignment = 1  # Center vertically
    shade_cell(merged_test46, 'E0E0E0')
    
    # Test 47: Uniformity (row 10)
    cell = table.cell(10, 0)
    cell.text = 'Test 47:'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    shade_cell(cell, 'E0E0E0')
    
    cell = table.cell(10, 1)
    cell.text = 'Uniformity:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(10, day + 1))
    
    # Image Noise
    cell = table.cell(11, 1)
    cell.text = 'Image Noise (± 10%) Mean:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(11, day + 1))
    
    # Standard Deviation
    cell = table.cell(12, 1)
    cell.text = 'Standard Deviation:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(12, day + 1))
    
    # Merge Test 47 cells
    merged_test47 = table.cell(10, 0).merge(table.cell(12, 0))
    merged_test47.text = 'Test 47:'
    merged_test47.paragraphs[0].runs[0].font.bold = True
    merged_test47.paragraphs[0].runs[0].font.size = Pt(10)
    merged_test47.vertical_alignment = 1
    shade_cell(merged_test47, 'E0E0E0')
    
    # Test 48 (row 13)
    cell = table.cell(13, 0)
    cell.text = 'Test 48:'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    shade_cell(cell, 'E0E0E0')
    
    cell = table.cell(13, 1)
    cell.text = 'H₂O HUs (± 5):'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(13, day + 1))
    
    # Add more rows for remaining items
    table.add_row()
    table.add_row()
    table.add_row()
    
    # Solid HUs
    cell = table.cell(14, 1)
    cell.text = 'Solid HUs (± 10):'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(14, day + 1))
    
    # Merge Test 48
    merged_test48 = table.cell(13, 0).merge(table.cell(14, 0))
    merged_test48.text = 'Test 48:'
    merged_test48.paragraphs[0].runs[0].font.bold = True
    merged_test48.paragraphs[0].runs[0].font.size = Pt(10)
    merged_test48.vertical_alignment = 1
    shade_cell(merged_test48, 'E0E0E0')
    
    # Test 49
    cell = table.cell(15, 0)
    cell.text = 'Test 49:'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    shade_cell(cell, 'E0E0E0')
    
    cell = table.cell(15, 1)
    cell.text = 'Scan Plane Localization from Alignment Lights:'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    shade_cell(cell, 'F9F9F9')
    for day in range(1, 32):
        add_checkbox(table.cell(15, day + 1))
    
    # Comments row
    cell = table.cell(16, 0)
    cell.text = 'COMMENTS'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    shade_cell(cell, 'E0E0E0')
    
    merged_comments = table.cell(16, 1).merge(table.cell(16, 32))
    
    # Add instructions
    doc.add_paragraph()
    instructions = doc.add_paragraph()
    instructions.add_run('Instructions: ').font.bold = True
    instructions_text = instructions.add_run(
        'Complete each test item daily and mark in the appropriate date column. '
        'Use ✓ for acceptable results, X for issues found, or record actual values where indicated. '
        'Add any relevant comments in the COMMENTS row.'
    )
    instructions_text.font.size = Pt(9)
    
    # Save document
    output_path = r'c:\Users\Admin\Desktop\ELC\Ubuntu-Patient-Care\Philips_QC_Checklist_Tests_46-49.docx'
    doc.save(output_path)
    print(f"Word document created successfully: {output_path}")
    return output_path

if __name__ == '__main__':
    try:
        create_philips_qc_checklist_2()
        print("\n✓ Document created with checkboxes and proper formatting!")
    except ImportError as e:
        print(f"Error: Missing required library. Please install python-docx:")
        print("pip install python-docx")
    except Exception as e:
        print(f"Error creating document: {e}")
        import traceback
        traceback.print_exc()
