"""
Script to create a Philips Radiography QC Checklist Word document
with proper tables and checkboxes
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_checkbox(cell):
    """Add a checkbox to a table cell"""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    # Add checkbox symbol
    run.text = '☐'
    run.font.size = Pt(14)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def shade_cell(cell, color):
    """Shade a cell with a specific color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def create_philips_qc_checklist():
    """Create the Philips QC Checklist Word document"""
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add Philips logo text (right-aligned)
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_run = logo_para.add_run('PHILIPS')
    logo_run.font.size = Pt(16)
    logo_run.font.bold = True
    logo_run.font.color.rgb = RGBColor(0, 0, 139)
    
    doc.add_paragraph()  # Space
    
    # Create table: 1 header column + 31 day columns = 32 columns
    # Rows: header info (4) + check items (19) = 23 rows
    table = doc.add_table(rows=23, cols=32)
    table.style = 'Table Grid'
    
    # Row 0: Month and Year
    cell = table.cell(0, 0)
    cell.text = 'Month and Year:'
    cell.paragraphs[0].runs[0].font.bold = True
    shade_cell(cell, 'F0F0F0')
    merged_cell = table.cell(0, 1).merge(table.cell(0, 31))
    shade_cell(merged_cell, 'FFFFFF')
    
    # Row 1: Date headers (1-31)
    date_cell = table.cell(1, 0)
    date_cell.text = 'Date:'
    date_cell.paragraphs[0].runs[0].font.bold = True
    shade_cell(date_cell, 'F0F0F0')
    
    for day in range(1, 32):
        cell = table.cell(1, day)
        cell.text = str(day)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(cell, 'F0F0F0')
    
    # Row 2: DoH Licence #
    cell = table.cell(2, 0)
    cell.text = 'DoH Licence #:'
    shade_cell(cell, 'F9F9F9')
    merged_cell = table.cell(2, 1).merge(table.cell(2, 31))
    
    # Row 3: Unit Serial #
    cell = table.cell(3, 0)
    cell.text = 'Unit Serial #:'
    shade_cell(cell, 'F9F9F9')
    merged_cell = table.cell(3, 1).merge(table.cell(3, 31))
    
    # Checklist items starting from row 4
    checklist_items = [
        'Radiographer Initials:',
        'Room Cleaned & Sponges Available:',
        'Tube Warm-Up:',
        'Lead Aprons:',
        'Radiation Warning Lights and Sign:',
        'Field Light:',
        'SID Marks:',
        'Panel Lights:',
        'Overhead Crane Movement:',
        'Door Movement:',
        'Unit Locks and Brakes:',
        'Centre Stop:',
        'Perpendicularity:',
        'Table Movement and Angulation:',
        'Skye Plate Functional:',
        'Transverse/Longitudinal Loose Grid:',
        'Cleanliness:',
        'Display Monitors - Flicker:',
        'Display Monitors - Distortion:',
    ]
    
    for idx, item in enumerate(checklist_items):
        row = 4 + idx
        cell = table.cell(row, 0)
        cell.text = item
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(cell, 'F9F9F9')
        
        # Add checkboxes for each day
        for day in range(1, 32):
            cell = table.cell(row, day)
            add_checkbox(cell)
    
    # Add instructions
    doc.add_paragraph()
    instructions = doc.add_paragraph()
    instructions.add_run('Instructions: ').font.bold = True
    instructions_text = instructions.add_run(
        'Check each item daily and mark in the appropriate date column. '
        'Use ✓ for acceptable, X for issues found, or leave blank if not applicable.'
    )
    instructions_text.font.size = Pt(9)
    
    # Save document
    output_path = r'c:\Users\Admin\Desktop\ELC\Ubuntu-Patient-Care\Philips_QC_Checklist.docx'
    doc.save(output_path)
    print(f"Word document created successfully: {output_path}")
    return output_path

if __name__ == '__main__':
    try:
        create_philips_qc_checklist()
        print("\n✓ Document created with checkboxes and proper formatting!")
    except ImportError as e:
        print(f"Error: Missing required library. Please install python-docx:")
        print("pip install python-docx")
    except Exception as e:
        print(f"Error creating document: {e}")
