#!/usr/bin/env python3
"""Convert PDF to editable Word document"""
import sys
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def convert_pdf_to_docx(pdf_path, output_path):
    """Extract text from PDF and create a Word document"""
    print(f"Reading PDF: {pdf_path}")
    
    # Read PDF
    reader = PdfReader(pdf_path)
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    print(f"Extracting text from {len(reader.pages)} pages...")
    
    # Extract text from each page
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        
        if text.strip():
            # Add page number as heading
            if page_num > 1:
                doc.add_page_break()
            
            # Add the text content
            lines = text.split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    # Check if line looks like a heading (all caps, short)
                    if len(line) < 100 and line.isupper() and len(line.split()) < 10:
                        heading = doc.add_heading(line, level=2)
                    else:
                        para = doc.add_paragraph(line)
    
    # Save document
    doc.save(output_path)
    print(f"✅ Conversion complete: {output_path}")
    print(f"You can now edit the document in Microsoft Word or any compatible editor")

if __name__ == '__main__':
    pdf_file = r"e:\Downloads\Business Application .pdf"
    output_file = r"e:\Downloads\Business Application - Editable.docx"
    
    try:
        convert_pdf_to_docx(pdf_file, output_file)
    except Exception as e:
        print(f"❌ Error: {e}")
        sys.exit(1)
