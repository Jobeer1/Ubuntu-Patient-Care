"""
Script to create a Philips QC Checklist Word document (DDR System, AEC Device, Image Display)
with proper tables and checkboxes
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_checkbox(cell):
    """Add a checkbox to a table cell"""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.text = '☐'
    run.font.size = Pt(14)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def shade_cell(cell, color):
    """Shade a cell with a specific color"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def set_cell_border(cell, **kwargs):
    """Set border properties for a cell"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            边 = OxmlElement(f'w:{edge}')
            边.set(qn('w:val'), 'single')
            边.set(qn('w:sz'), '12')
            边.set(qn('w:color'), '000000')
            tcBorders.append(边)
    
    tcPr.append(tcBorders)

def create_philips_qc_checklist_3():
    """Create the Philips QC Checklist Word document (DDR/AEC/Display Tests)"""
    doc = Document()
    
    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.page_width = Inches(11)  # Landscape
        section.page_height = Inches(8.5)
    
    # Add Philips logo text (right-aligned)
    logo_para = doc.add_paragraph()
    logo_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    logo_run = logo_para.add_run('PHILIPS')
    logo_run.font.size = Pt(16)
    logo_run.font.bold = True
    logo_run.font.color.rgb = RGBColor(0, 0, 139)
    
    doc.add_paragraph()  # Space
    
    # Create main table
    # Structure: Test ID | Category | Frequency | Baseline | Results | 6 checkbox columns
    table = doc.add_table(rows=15, cols=11)
    table.style = 'Table Grid'
    
    # Set column widths
    table.columns[0].width = Inches(0.6)   # Test ID
    table.columns[1].width = Inches(1.5)   # Category/Test Name
    table.columns[2].width = Inches(0.8)   # Frequency
    table.columns[3].width = Inches(2.5)   # Baseline
    table.columns[4].width = Inches(2.5)   # Results
    for i in range(5, 11):
        table.columns[i].width = Inches(0.5)  # Checkbox columns
    
    row_idx = 0
    
    # ===== III.1.6 DDR SYSTEM =====
    # Header Row
    cell = table.cell(row_idx, 0)
    cell.text = 'III.1.6'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'D3D3D3')
    
    cell = table.cell(row_idx, 1)
    cell.text = 'DDR SYSTEM'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'D3D3D3')
    
    cell = table.cell(row_idx, 3)
    cell.text = 'Baseline'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'D3D3D3')
    
    cell = table.cell(row_idx, 4)
    cell.text = 'Results'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'D3D3D3')
    
    # Add checkboxes to header
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    row_idx += 1
    
    # Test 24: Detector dose indicator monitoring
    cell = table.cell(row_idx, 0)
    cell.text = '24'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'E8E8E8')
    
    cell = table.cell(row_idx, 1)
    cell.text = 'Detector dose indicator monitoring'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    cell = table.cell(row_idx, 2)
    cell.text = '3 Monthly'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Baseline column
    cell = table.cell(row_idx, 3)
    baseline_para = cell.paragraphs[0]
    baseline_para.add_run('Exposure Index:\n').font.size = Pt(8)
    baseline_para.add_run('• 20½ = _________\n').font.size = Pt(8)
    baseline_para.add_run('• 20v = _________').font.size = Pt(8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Results column
    cell = table.cell(row_idx, 4)
    result_para = cell.paragraphs[0]
    result_para.add_run('Exposure Index: _________').font.size = Pt(8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Add checkboxes
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    row_idx += 1
    
    # Test 25: Image uniformity
    cell = table.cell(row_idx, 0)
    cell.text = '25'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'E8E8E8')
    
    cell = table.cell(row_idx, 1)
    cell.text = 'Image uniformity'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    cell = table.cell(row_idx, 2)
    cell.text = '3 monthly'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Add checkboxes
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    row_idx += 1
    
    # ===== 111.1.6.1 AEC DEVICE =====
    cell = table.cell(row_idx, 0)
    cell.text = '111.1.6.1'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'D3D3D3')
    
    cell = table.cell(row_idx, 1)
    cell.text = 'AEC DEVICE'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'D3D3D3')
    
    # Add checkboxes
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    row_idx += 1
    
    # Test 27: Sensitivity (test all chambers)
    # This is a complex multi-row section
    # First row with test number
    start_row = row_idx
    
    cell = table.cell(row_idx, 0)
    cell.text = '27'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'E8E8E8')
    
    cell = table.cell(row_idx, 1)
    cell.text = 'Sensitivity (test all chambers)'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    cell = table.cell(row_idx, 2)
    cell.text = '3 Monthly'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Baseline - Left Chamber
    cell = table.cell(row_idx, 3)
    para = cell.paragraphs[0]
    para.add_run('Left Chamber\n').font.bold = True
    para.runs[0].font.size = Pt(8)
    para.add_run('mAs = _______\n').font.size = Pt(7)
    para.add_run('±±5% = _______\n').font.size = Pt(7)
    para.add_run('Cl = _______\n').font.size = Pt(7)
    para.add_run('±±5% = _______').font.size = Pt(7)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Results - Left Chamber
    cell = table.cell(row_idx, 4)
    para = cell.paragraphs[0]
    para.add_run('Left Chamber\n').font.bold = True
    para.runs[0].font.size = Pt(8)
    para.add_run('mAs = _______\n').font.size = Pt(7)
    para.add_run('Cl = _______').font.size = Pt(7)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Add checkboxes
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    row_idx += 1
    
    # Middle Chamber
    cell = table.cell(row_idx, 3)
    para = cell.paragraphs[0]
    para.add_run('Middle Chamber\n').font.bold = True
    para.runs[0].font.size = Pt(8)
    para.add_run('mAs = _______\n').font.size = Pt(7)
    para.add_run('±±5% = _______\n').font.size = Pt(7)
    para.add_run('Cl = _______\n').font.size = Pt(7)
    para.add_run('±±5% = _______').font.size = Pt(7)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    cell = table.cell(row_idx, 4)
    para = cell.paragraphs[0]
    para.add_run('Middle Chamber\n').font.bold = True
    para.runs[0].font.size = Pt(8)
    para.add_run('mAs = _______\n').font.size = Pt(7)
    para.add_run('Cl = _______').font.size = Pt(7)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Add checkboxes
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    row_idx += 1
    
    # Right Chamber
    cell = table.cell(row_idx, 3)
    para = cell.paragraphs[0]
    para.add_run('Right Chamber\n').font.bold = True
    para.runs[0].font.size = Pt(8)
    para.add_run('mAs = _______\n').font.size = Pt(7)
    para.add_run('±±5% = _______\n').font.size = Pt(7)
    para.add_run('Cl = _______\n').font.size = Pt(7)
    para.add_run('±±5% = _______').font.size = Pt(7)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    cell = table.cell(row_idx, 4)
    para = cell.paragraphs[0]
    para.add_run('Right Chamber\n').font.bold = True
    para.runs[0].font.size = Pt(8)
    para.add_run('mAs = _______\n').font.size = Pt(7)
    para.add_run('Cl = _______').font.size = Pt(7)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Add checkboxes
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    # Merge cells for test 27
    end_row = row_idx
    merged = table.cell(start_row, 0).merge(table.cell(end_row, 0))
    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    merged = table.cell(start_row, 1).merge(table.cell(end_row, 1))
    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    merged = table.cell(start_row, 2).merge(table.cell(end_row, 2))
    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    row_idx += 1
    
    # ===== 111.1.8 IMAGE DISPLAY =====
    cell = table.cell(row_idx, 0)
    cell.text = '111.1.8'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'D3D3D3')
    
    cell = table.cell(row_idx, 1)
    cell.text = 'IMAGE DISPLAY'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(11)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'D3D3D3')
    
    # Add checkboxes
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    row_idx += 1
    
    # Test 29a: Conditions of image display
    start_row_29 = row_idx
    
    cell = table.cell(row_idx, 0)
    cell.text = '29 a.'
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade_cell(cell, 'E8E8E8')
    
    cell = table.cell(row_idx, 1)
    cell.text = 'Conditions of image display'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    cell = table.cell(row_idx, 2)
    cell.text = '6 MONTHLY'
    cell.paragraphs[0].runs[0].font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Results column with questions
    cell = table.cell(row_idx, 4)
    para = cell.paragraphs[0]
    para.add_run('Monitors clean and without damage?    YES/NO').font.size = Pt(8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Add checkboxes
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    row_idx += 1
    
    # Question 2
    cell = table.cell(row_idx, 4)
    para = cell.paragraphs[0]
    para.add_run('Monitors free from flicker and visibly clear? YES/NO').font.size = Pt(8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    row_idx += 1
    
    # Question 3
    cell = table.cell(row_idx, 4)
    para = cell.paragraphs[0]
    para.add_run('Monitors compare well with each other?    YES/NO').font.size = Pt(8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    row_idx += 1
    
    # Question 4
    cell = table.cell(row_idx, 4)
    para = cell.paragraphs[0]
    para.add_run('Distortions or non-uniformities?    YES/NO').font.size = Pt(8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    for col in range(5, 11):
        add_checkbox(table.cell(row_idx, col))
    
    # Merge cells for test 29a
    end_row_29 = row_idx
    merged = table.cell(start_row_29, 0).merge(table.cell(end_row_29, 0))
    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    merged = table.cell(start_row_29, 1).merge(table.cell(end_row_29, 1))
    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    merged = table.cell(start_row_29, 2).merge(table.cell(end_row_29, 2))
    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    
    # Add instructions
    doc.add_paragraph()
    instructions = doc.add_paragraph()
    instructions.add_run('Instructions: ').font.bold = True
    instructions_text = instructions.add_run(
        'Complete each test at the specified frequency. Record baseline values and compare results. '
        'Mark checkboxes for each test completion. Answer YES/NO questions as appropriate. '
        'Ensure all measurements are within acceptable tolerances (±5% where indicated).'
    )
    instructions_text.font.size = Pt(9)
    
    # Save document
    output_path = r'c:\Users\Admin\Desktop\ELC\Ubuntu-Patient-Care\Philips_QC_Checklist_DDR_AEC_Display.docx'
    doc.save(output_path)
    print(f"Word document created successfully: {output_path}")
    return output_path

if __name__ == '__main__':
    try:
        create_philips_qc_checklist_3()
        print("\n✓ Document created with checkboxes and proper formatting!")
    except ImportError as e:
        print(f"Error: Missing required library. Please install python-docx:")
        print("pip install python-docx")
    except Exception as e:
        print(f"Error creating document: {e}")
        import traceback
        traceback.print_exc()
